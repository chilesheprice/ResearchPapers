#!/usr/bin/env python3
"""
Convert the Menopause and the Heart 2026 Update markdown to Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document(md_file, output_file):
    """Convert markdown to Word document with proper formatting"""

    # Create a new Document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip horizontal rules
        if line == '---':
            i += 1
            continue

        # Handle headings
        if line.startswith('# '):
            # Main title (H1)
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif line.startswith('## '):
            # Section heading (H2)
            doc.add_heading(line[3:], level=2)

        elif line.startswith('### '):
            # Subsection heading (H3)
            doc.add_heading(line[4:], level=3)

        elif line.startswith('#### '):
            # Sub-subsection heading (H4)
            doc.add_heading(line[5:], level=4)

        # Handle bold text (author name, etc.)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('*'))
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Handle bullet points
        elif line.startswith('- ') or line.startswith('• '):
            # Remove the bullet/dash
            text = line[2:] if line.startswith('- ') else line[2:]

            # Check indentation level
            indent_level = 0
            if lines[i].startswith('  '):
                indent_level = 1

            p = doc.add_paragraph(text, style='List Bullet')

        # Handle numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(text, style='List Number')

        # Handle regular paragraphs
        elif line:
            # Check if it's italicized
            if line.startswith('*') and line.endswith('*') and not line.startswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(line.strip('*'))
                run.italic = True
            else:
                # Regular paragraph - handle inline formatting
                p = doc.add_paragraph()

                # Split by bold markers
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part.strip('*'))
                        run.bold = True
                    elif part:
                        p.add_run(part)

        i += 1

    # Add page breaks before major sections if needed
    # This can be customized based on requirements

    # Save the document
    doc.save(output_file)
    print(f"Word document created successfully: {output_file}")

if __name__ == "__main__":
    md_file = "/Users/chileshe/Desktop/ResearchPapers/Menopause_and_the_Heart_2026_Update.md"
    output_file = "/Users/chileshe/Desktop/ResearchPapers/Menopause_and_the_Heart_2026_Update.docx"

    create_word_document(md_file, output_file)
